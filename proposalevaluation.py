import streamlit as st
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
import google.generativeai as genai

genai.configure(api_key=st.secrets["google_gen_ai"]["api_key"])

def read_pdf(uploaded_file):
    file_stream = uploaded_file.read()
    doc = fitz.open("pdf", file_stream)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def evaluate_with_gemini(proposal_text, sections, expertise):
    model = genai.GenerativeModel("gemini-pro")
    chat = model.start_chat(history=[])
    responses = {}
    total_max_points = 0
    total_awarded_points = 0
    for section in sections:
        prompt = f"""
        You're an expert in {expertise}, known for two decades of meticulous study, review, and evaluation of {expertise} projects and proposals.

        You will be given a proposal submission to consider. Your task will involve a comprehensive review of it based on your subject matter expertise, and the project's defined sections.

        In your evaluation, you will provide an overall score supported by detailed commentary on each section. Additionally, you will write a report offering feedback to the Respondent and suggestions for improving the submission. Areas to focus on include:

        Prose: Assess the clarity and precision of the language used in each section. Evaluate how effectively the language facilitates comprehension of the offerings and methodologies.

        Structure: Review the logical flow and coherence of the sections. Ensure that the points are presented in a well-organized manner that aligns with procurement evaluation protocols.

        Format: Evaluate the professional formatting of the documents. Consider how the formatting enhances readability, making it easier for procurement teams to locate critical information.

        Value: Whenever prices or fees are presented, please assess the extent to which the dollar values seem to be a good value in terms of the services offered.
        """
        response = chat.send_message(prompt)
        evaluation = response.text.strip()
        max_points = int(section['points'])
        total_max_points += max_points
        score = calculate_score(evaluation, max_points)
        awarded_points = score / max_points * 100  # Calculate awarded percentage
        responses[section['name']] = {
            'evaluation': evaluation,
            'awarded_percentage': awarded_points,  # Store as percentage
            'max_points': max_points
        }
        total_awarded_points += score * max_points / 100  # Convert percentage back for total calculation
    
    overall_score = total_awarded_points / total_max_points if total_max_points > 0 else 0
    
    return responses, overall_score

def calculate_score(evaluation_text, max_points):
    max_points = int(max_points)
    evaluation_quality = evaluate_quality(evaluation_text)
    if evaluation_quality == "excellent":
        score = int(0.9 * max_points)
    elif evaluation_quality == "good":
        score = int(0.8 * max_points)
    elif evaluation_quality == "average":
        score = int(0.7 * max_points)
    elif evaluation_quality == "poor":
        score = int(0.6 * max_points)
    elif evaluation_quality == "very poor":
        score = int(0.5 * max_points)
    else:
        score = 0
    return score

def evaluate_quality(evaluation_text):
    if "relevant" in evaluation_text.lower() and "well-analyzed" in evaluation_text.lower():
        return "excellent"
    elif "relevant" in evaluation_text.lower():
        return "good"
    elif "clear" in evaluation_text.lower():
        return "average"
    elif "unclear" in evaluation_text.lower():
        return "poor"
    else:
        return "very poor"

def create_docx(report_data, overall_score):
    doc = Document()
    doc.add_heading("Proposal Evaluation Report", level=1)
    for section, data in report_data.items():
        doc.add_heading(f"Section: {section}", level=2)
        doc.add_paragraph(f"Evaluation: {data['evaluation']}")
        doc.add_paragraph(f"Score: {data['awarded_percentage']:.1f}% ({data['max_points']} points maximum)")
    doc.add_heading("Overall Score", level=2)
    doc.add_paragraph(f"Overall Score: {overall_score:.1%}")
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output

def display_initial_evaluations(evaluations):
    for section, data in evaluations.items():
        st.write(f"### {section}")
        st.text_area(f"Evaluation for '{section}'", value=data['evaluation'], key=f"eval_{section}_view")
        st.write("Score:", f"{data['awarded_percentage']:.1f}%", f"({data['max_points']} points maximum)")

def display_revision_interface(evaluations):
    for section, data in evaluations.items():
        with st.container():
            eval_key = f"eval_{section}_edit"
            data['evaluation'] = st.text_area(f"Edit evaluation for '{section}':", value=data['evaluation'], key=eval_key)
            max_points_key = f"max_points_{section}_edit"
            data['max_points'] = st.text_input(f"Max points for '{section}':", value=data['max_points'], key=max_points_key)

def main():
    st.title("Proposal Evaluation App")
    expertise = st.text_input("Enter your field of expertise", value="environmental science")
    num_sections = st.number_input("How many sections does your proposal have?", min_value=1, max_value=10, step=1)
    sections = []
    for i in range(int(num_sections)):
        with st.expander(f"Section {i + 1} Details"):
            section_name = st.text_input(f"Name of section {i + 1}")
            section_points = st.text_input(f"Max points for '{section_name}'", value="5", key=f"max_points_{i}")
            sections.append({'name': section_name, 'points': int(section_points)})

    uploaded_file = st.file_uploader("Upload your proposal PDF", type=["pdf"])
    if uploaded_file is not None:
        proposal_text = read_pdf(uploaded_file)
        evaluations, overall_score = evaluate_with_gemini(proposal_text, sections, expertise)
        st.session_state.evaluations = evaluations  # Initialize or update the session state with evaluations
        display_initial_evaluations(evaluations)

        if st.button("Review and Edit Evaluations"):
            display_revision_interface(st.session_state.evaluations)
            if st.button("Save Edited Evaluations"):
                st.experimental_rerun()  # Rerun the app to reflect the updated evaluations

        if st.button("Finalize and Download Report"):
            evaluations = st.session_state.evaluations
            docx_file = create_docx(evaluations, overall_score)
            st.download_button(
                label="Download Evaluation Report",
                data=docx_file,
                file_name="evaluation_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
